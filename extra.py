
1. yes becous index 5 doeswnt exist.

2.
colors = ["red", "green", "blue", "purple"]
for i in range(len(colors)):
    print(colors[2])

3.
creating a file on the main disk

4.
protect the code by puting mode "a" that open for appending the file without truncating it.
creates a new file if it does not exist.

5.
import os
directory = "test"
parent_dir = "/Users/bellamarkovitz/Documents/devops/"
path = os.path.join(parent_dir, directory)
os.mkdir(path)
print("Directory '% s' created" % directory)

my_file = open("/Users/bellamarkovitz/Documents/devops/test/name.txt","a",encoding="utf-8")
content = my_file.write("Bella Markovitz")

6.
from docx import Document
document = Document()
document.save("test.docx")

exstra
import docx
doc = docx.Document()
doc.add_heading()
doc_para = doc.add_paragraph("Bella Markovitz")
doc.save("test.docx")